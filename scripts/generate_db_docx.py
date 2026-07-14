import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

def add_page_number(run):
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    
    r = run._r
    r.append(fldChar1)
    r.append(instrText)
    r.append(fldChar2)
    r.append(fldChar3)

def style_paragraph(p, font_name="Times New Roman", size_pt=13, line_spacing=1.5, space_before_pt=0, space_after_pt=6, bold=False, italic=False, align_justify=True):
    p.paragraph_format.line_spacing = line_spacing
    p.paragraph_format.space_before = Pt(space_before_pt)
    p.paragraph_format.space_after = Pt(space_after_pt)
    if align_justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    for r in p.runs:
        r.font.name = font_name
        r.font.size = Pt(size_pt)
        r.bold = bold
        r.italic = italic
        r.font.color.rgb = RGBColor(0, 0, 0)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def set_table_borders(table):
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>\n'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>\n'
        f'  <w:bottom w:val="single" w:sz="6" w:space="0" w:color="333333"/>\n'
        f'  <w:left w:val="none"/>\n'
        f'  <w:right w:val="none"/>\n'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="E0E0E0"/>\n'
        f'  <w:insideV w:val="none"/>\n'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

def add_styled_table(doc, headers, data):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    
    # Header Row
    hdr_cells = table.rows[0].cells
    for i, title in enumerate(headers):
        hdr_cells[i].text = title
        set_cell_margins(hdr_cells[i], top=120, bottom=120, left=150, right=150)
        set_cell_shading(hdr_cells[i], '4FAF4F') # Green steps green
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for r in p.runs:
            r.font.name = "Times New Roman"
            r.font.size = Pt(11)
            r.bold = True
            r.font.color.rgb = RGBColor(255, 255, 255) # White text
            
    # Data Rows
    for row_idx, row_data in enumerate(data):
        row_cells = table.add_row().cells
        for col_idx, text in enumerate(row_data):
            row_cells[col_idx].text = str(text)
            set_cell_margins(row_cells[col_idx], top=100, bottom=100, left=150, right=150)
            if row_idx % 2 == 1:
                set_cell_shading(row_cells[col_idx], 'F9FBF9') # Zebra
            p = row_cells[col_idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for r in p.runs:
                r.font.name = "Times New Roman"
                r.font.size = Pt(10.5)
                r.font.color.rgb = RGBColor(51, 51, 51)
                
    # Add spacing after table
    p_spacer = doc.add_paragraph()
    style_paragraph(p_spacer, size_pt=6, space_after_pt=6)

def main():
    doc = Document()
    
    # Page Setup
    for section in doc.sections:
        section.top_margin = Inches(0.79)
        section.bottom_margin = Inches(0.79)
        section.left_margin = Inches(1.18)
        section.right_margin = Inches(0.79)
        
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        footer_para.text = ""
        run = footer_para.add_run("Đặc tả Cơ sở Dữ liệu GreenSteps | Trang ")
        run.font.name = "Times New Roman"
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(128, 128, 128)
        
        page_run = footer_para.add_run()
        page_run.font.name = "Times New Roman"
        page_run.font.size = Pt(10)
        page_run.font.color.rgb = RGBColor(128, 128, 128)
        add_page_number(page_run)

    # Document Header
    p_title = doc.add_paragraph()
    p_title.add_run("TÀI LIỆU ĐẶC TẢ CƠ SỞ DỮ LIỆU").bold = True
    style_paragraph(p_title, size_pt=18, bold=True, space_after_pt=6, align_justify=False)
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p_subtitle = doc.add_paragraph()
    p_subtitle.add_run("Hệ Thống Du Lịch Xanh GreenSteps (PostgreSQL)")
    style_paragraph(p_subtitle, size_pt=14, italic=True, space_after_pt=24, align_justify=False)
    p_subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p_intro = doc.add_paragraph()
    p_intro.add_run("Tài liệu này đặc tả chi tiết cấu trúc các bảng dữ liệu trong hệ thống cơ sở dữ liệu quan hệ PostgreSQL của dự án GreenSteps. Các bảng được quản lý tập trung thông qua Sequelize ORM, hỗ trợ đầy đủ các mối quan hệ khóa ngoại cứng, tính toàn vẹn dữ liệu toàn hệ thống, các ràng buộc duy nhất (Unique), chỉ mục (Indexes), và đặc biệt đảm bảo tuyệt đối tính chất giao dịch ACID phục vụ các nghiệp vụ ví tiền điện tử, ký quỹ e-voucher và thanh toán đặt chỗ trực tuyến.")
    style_paragraph(p_intro)

    # Tables specs in 5 columns format
    tables_data = [
        {
            "name": "1. Badge (Nhãn dán)",
            "desc": "Lưu trữ các nhãn dán xanh hoặc danh hiệu áp dụng cho người dùng hoặc lịch trình/dịch vụ.",
            "fields": [
                ["ID_Badge", "VARCHAR(50)", "PRIMARY KEY", "Mã nhãn dán", "Tên nhãn dán viết liền, duy nhất"],
                ["badges_description", "VARCHAR(255)", "NULL", "Mô tả ý nghĩa của nhãn dán", ""],
                ["foruserortour", "INT", "NOT NULL", "Phân loại nhãn dán", "0 ứng với người dùng, 1 ứng với tour/dịch vụ"]
            ]
        },
        {
            "name": "2. User (Người dùng)",
            "desc": "Lưu trữ tài khoản của tất cả người dùng trong hệ thống (Traveler, Provider, Admin).",
            "fields": [
                ["ID_User", "VARCHAR(50)", "PRIMARY KEY", "Mã người dùng", "UG + 2 số cuối năm + 3 chữ đầu username + số thứ tự (0000-9999)"],
                ["role", "VARCHAR(10)", "NOT NULL", "Vai trò người dùng", "0 ứng với traveler, 1 ứng với provider, 2 ứng với admin"],
                ["username", "VARCHAR(50)", "UNIQUE, NOT NULL", "Tên đăng nhập", "Chữ cái và số, tối thiểu 3 kí tự, tối đa 20 kí tự"],
                ["password", "VARCHAR(255)", "NOT NULL", "Mật khẩu đăng nhập", "Chuỗi ký tự đã băm mã hóa"],
                ["fullname", "VARCHAR(100)", "NOT NULL", "Họ và tên đầy đủ", ""],
                ["email", "VARCHAR(100)", "UNIQUE, NOT NULL", "Địa chỉ thư điện tử", ""],
                ["phone", "VARCHAR(20)", "NULL", "Số điện thoại liên lạc", ""],
                ["dob", "VARCHAR(20)", "NULL", "Ngày sinh", "Định dạng DD/MM/YYYY"],
                ["gender", "CHAR(1)", "NULL", "Giới tính", "0 là nữ, 1 là nam, 2 là khác"],
                ["address", "VARCHAR(255)", "NULL", "Địa chỉ liên lạc", "Địa chỉ chi tiết (2 kiểu dữ liệu: address_line_1/ward/province/country hoặc ward/province/country)"],
                ["job", "VARCHAR(100)", "NULL", "Nghề nghiệp", ""],
                ["company_name", "VARCHAR(100)", "NULL", "Tên công ty doanh nghiệp đối tác", "Chỉ áp dụng với tài khoản đối tác (role = 1)"]
            ]
        },
        {
            "name": "3. BadgeUser (Nhãn dán Người dùng)",
            "desc": "Bảng liên kết thể hiện các danh hiệu/nhãn dán mà một người dùng đã đạt được.",
            "fields": [
                ["ID_Badge", "VARCHAR(50)", "PRIMARY KEY, FOREIGN KEY", "Mã nhãn dán liên kết", "Tham chiếu tới Badge(ID_Badge)"],
                ["ID_User", "VARCHAR(50)", "PRIMARY KEY, FOREIGN KEY", "Mã người dùng được gắn nhãn", "Tham chiếu tới User(ID_User)"]
            ]
        },
        {
            "name": "4. Wallet (Ví điện tử)",
            "desc": "Lưu trữ số dư và trạng thái ví điện tử tích hợp của người dùng để thực hiện thanh toán.",
            "fields": [
                ["ID_Wallet", "VARCHAR(50)", "PRIMARY KEY", "Mã ví điện tử", "EW + 8 số ngẫu nhiên"],
                ["ID_User", "VARCHAR(50)", "UNIQUE, FOREIGN KEY", "Mã người dùng sở hữu ví", "Tham chiếu tới User(ID_User)"],
                ["balance", "DECIMAL(12,2)", "NOT NULL", "Số dư ví hiện tại", "Mặc định 0.00, không âm"],
                ["green_points", "INT", "NOT NULL", "Điểm xanh tích lũy", "Mặc định 0, tích lũy dựa trên hoạt động giảm thiểu carbon"],
                ["registered", "BOOLEAN", "NOT NULL", "Trạng thái kích hoạt ví", "True: đã kích hoạt, False: chưa kích hoạt"]
            ]
        },
        {
            "name": "5. WalletTransaction (Giao dịch ví)",
            "desc": "Lưu lịch sử nạp tiền, thanh toán và hoàn tiền qua ví điện tử.",
            "fields": [
                ["ID_Transaction", "VARCHAR(50)", "PRIMARY KEY", "Mã định danh giao dịch", "GD + timestamp"],
                ["ID_Wallet", "VARCHAR(50)", "FOREIGN KEY", "Mã ví thực hiện giao dịch", "Tham chiếu tới Wallet(ID_Wallet)"],
                ["type", "VARCHAR(20)", "NOT NULL", "Loại giao dịch", "deposit (nạp tiền), payment (thanh toán), refund (hoàn tiền)"],
                ["description", "VARCHAR(255)", "NULL", "Mô tả nội dung giao dịch", ""],
                ["amount", "DECIMAL(12,2)", "NOT NULL", "Số tiền giao dịch", "Giá trị dương cho nạp/hoàn tiền, âm cho thanh toán"],
                ["status", "VARCHAR(20)", "NOT NULL", "Trạng thái giao dịch", "success (thành công), pending (đang chờ), failed (thất bại)"],
                ["reference_id", "VARCHAR(100)", "NULL", "Mã tham chiếu thanh toán", "Lưu mã giao dịch ngân hàng đối tác hoặc thông báo đối soát"]
            ]
        },
        {
            "name": "6. Vender (Đối tác Nhà cung cấp)",
            "desc": "Thông tin bổ sung của các tài khoản có vai trò nhà cung ứng dịch vụ (Provider).",
            "fields": [
                ["ID_Vender", "VARCHAR(50)", "PRIMARY KEY", "Mã định danh đối tác", "vender_ + 6 số cuối timestamp"],
                ["ID_User", "VARCHAR(50)", "UNIQUE, FOREIGN KEY", "Mã người dùng liên kết", "Tham chiếu tới User(ID_User)"],
                ["registration_date", "DATE", "NOT NULL", "Ngày đăng ký đối tác", "Định dạng YYYY-MM-DD"]
            ]
        },
        {
            "name": "7. VenderContract (Hợp đồng Đối tác)",
            "desc": "Lưu trữ các văn bản hợp đồng pháp lý được ký giữa hệ thống và nhà cung cấp dịch vụ.",
            "fields": [
                ["ID_VenderContract", "VARCHAR(50)", "PRIMARY KEY", "Mã hợp đồng đối tác", "VC- + 6 số cuối timestamp"],
                ["ID_User", "VARCHAR(50)", "FOREIGN KEY", "Mã người dùng của nhà cung cấp", "Tham chiếu tới User(ID_User)"],
                ["name_contract", "VARCHAR(100)", "NOT NULL", "Tên loại hợp đồng dịch vụ", ""],
                ["text", "TEXT", "NOT NULL", "Toàn văn điều khoản hợp đồng", "Cam kết chiết khấu 10% và tiêu chuẩn du lịch xanh"]
            ]
        },
        {
            "name": "8. Revenue (Báo cáo doanh thu)",
            "desc": "Bảng tổng hợp đối soát tài chính định kỳ theo tháng của các nhà cung cấp.",
            "fields": [
                ["ID_Revenue", "VARCHAR(50)", "PRIMARY KEY", "Mã báo cáo doanh thu", "REV + 6 số cuối timestamp"],
                ["ID_User", "VARCHAR(50)", "FOREIGN KEY", "Mã nhà cung cấp được đối soát", "Tham chiếu tới User(ID_User)"],
                ["monthyear", "VARCHAR(10)", "NOT NULL", "Thời điểm chốt doanh thu", "Định dạng MM/YYYY"],
                ["total_booking", "INT", "NOT NULL", "Tổng số đơn đặt dịch vụ thành công", ""],
                ["total_revenue", "DECIMAL(12,2)", "NOT NULL", "Tổng doanh thu tích lũy", ""],
                ["service_fee", "DECIMAL(12,2)", "NOT NULL", "Phí dịch vụ trích lại cho GreenSteps", "10% tổng doanh thu"],
                ["final_profit", "DECIMAL(12,2)", "NOT NULL", "Lợi nhuận thực tế đối tác nhận", "90% tổng doanh thu"]
            ]
        },
        {
            "name": "9. Contract (Mẫu hợp đồng hệ thống)",
            "desc": "Mẫu hợp đồng quy chuẩn và trạng thái áp dụng trên toàn hệ thống.",
            "fields": [
                ["ID_Contract", "VARCHAR(50)", "PRIMARY KEY", "Mã mẫu hợp đồng", "CON + số thứ tự (ví dụ: CON0001)"],
                ["name_contract", "VARCHAR(100)", "NOT NULL", "Tên văn bản mẫu hợp đồng", ""],
                ["text", "TEXT", "NOT NULL", "Nội dung điều khoản mẫu", ""],
                ["contract_status", "VARCHAR(20)", "NOT NULL", "Trạng thái mẫu hợp đồng", "active (áp dụng) hoặc inactive (ngừng áp dụng)"]
            ]
        },
        {
            "name": "10. Provider (Công ty cung ứng dịch vụ)",
            "desc": "Hồ sơ thông tin chi nhánh/đơn vị cung ứng liên kết với các mẫu hợp đồng hệ thống.",
            "fields": [
                ["ID_Provider", "VARCHAR(50)", "PRIMARY KEY", "Mã định danh đơn vị", "PROV- + 6 số cuối của timestamp"],
                ["ID_Contract", "VARCHAR(50)", "FOREIGN KEY", "Mã hợp đồng mẫu áp dụng", "Tham chiếu tới Contract(ID_Contract)"],
                ["name_provider", "VARCHAR(100)", "NOT NULL", "Tên doanh nghiệp cung ứng", "Tên duy nhất thể hiện thương hiệu"],
                ["field", "VARCHAR(100)", "NOT NULL", "Lĩnh vực dịch vụ hoạt động", "Ví dụ: Dịch vụ lữ hành, dịch vụ lưu trú sinh thái"],
                ["destination", "VARCHAR(100)", "NOT NULL", "Khu vực/Địa điểm hoạt động", "Tỉnh thành hoạt động chính"],
                ["image_url", "VARCHAR(255)", "NULL", "Đường dẫn ảnh đại diện", ""],
                ["provider_status", "VARCHAR(20)", "NOT NULL", "Trạng thái hoạt động đối tác", "pending (chờ duyệt), active (đang hoạt động)"]
            ]
        },
        {
            "name": "11. Schedule (Lịch trình gốc / Khung hành trình)",
            "desc": "Khung thông tin chung của một chuyến đi bao gồm tên, điểm đến, số ngày và phát thải carbon.",
            "fields": [
                ["ID_Schedule", "VARCHAR(50)", "PRIMARY KEY", "Mã định danh lịch trình", "SCH + số thứ tự hoặc iti + số ngẫu nhiên"],
                ["tour_name", "VARCHAR(100)", "NOT NULL", "Tên chương trình du lịch", "Đại diện hành trình chính"],
                ["destination", "VARCHAR(100)", "NOT NULL", "Tỉnh thành/Điểm đến chính", "Nơi diễn ra chuyến đi"],
                ["days", "INT", "NOT NULL", "Tổng số ngày của chuyến đi", "Tối thiểu 1 ngày"],
                ["discount", "DECIMAL(5,2)", "NOT NULL", "Tỷ lệ giảm giá áp dụng", "Mặc định 0.00, biểu thị tỷ lệ phần trăm (0 - 100)"],
                ["carbon", "DECIMAL(10,2)", "NOT NULL", "Lượng carbon giảm thiểu hoặc phát thải", "Đơn vị kg CO2, giá trị âm thể hiện giảm thiểu xanh"],
                ["image_url", "VARCHAR(255)", "NULL", "Đường dẫn ảnh nền hành trình", ""],
                ["tour_description", "TEXT", "NULL", "Mô tả khái quát hành trình", ""],
                ["votes_count", "INT", "NOT NULL", "Số lượt bình chọn lịch trình", "Mặc định 0"]
            ]
        },
        {
            "name": "12. BadgeSchedule (Nhãn dán lịch trình)",
            "desc": "Bảng liên kết thể hiện các nhãn dán xanh được gán cho một lịch trình cụ thể.",
            "fields": [
                ["ID_Badge", "VARCHAR(50)", "PRIMARY KEY, FOREIGN KEY", "Mã nhãn dán liên kết", "Tham chiếu tới Badge(ID_Badge)"],
                ["ID_Schedule", "VARCHAR(50)", "PRIMARY KEY, FOREIGN KEY", "Mã lịch trình liên kết", "Tham chiếu tới Schedule(ID_Schedule)"]
            ]
        },
        {
            "name": "13. ScheduleSample (Hành trình mẫu có sẵn - Preset Tour)",
            "desc": "Các gói tour mẫu do các đối tác thiết lập có giá cả và điểm đánh giá cụ thể.",
            "fields": [
                ["ID_Schedule", "VARCHAR(50)", "PRIMARY KEY, FOREIGN KEY", "Mã hành trình", "Ánh xạ 1-1 tới Schedule(ID_Schedule)"],
                ["ID_Provider", "VARCHAR(50)", "FOREIGN KEY", "Đơn vị cung ứng/đối tác đại diện", "Tham chiếu tới Provider(ID_Provider)"],
                ["cost", "DECIMAL(12,2)", "NOT NULL", "Giá tiền trọn gói thực tế của tour", "Đơn vị VNĐ, không âm"],
                ["old_cost", "DECIMAL(12,2)", "NULL", "Giá gốc trước khi giảm giá", "Đơn vị VNĐ, dùng để hiển thị giá cũ gạch ngang"],
                ["rating", "DECIMAL(3,2)", "NOT NULL", "Điểm đánh giá trung bình", "Mặc định 5.0 (thang điểm 1 - 5 sao)"],
                ["votes_count", "INT", "NOT NULL", "Tổng số lượt đánh giá từ khách", "Mặc định 0, tự động cộng dồn qua hooks"]
            ]
        },
        {
            "name": "14. ScheduleCustom (Hành trình tự tùy biến)",
            "desc": "Hành trình do khách du lịch tự lên kế hoạch và chỉnh sửa chi phí.",
            "fields": [
                ["ID_Schedule", "VARCHAR(50)", "PRIMARY KEY, FOREIGN KEY", "Mã hành trình", "Ánh xạ 1-1 tới Schedule(ID_Schedule)"],
                ["ID_User", "VARCHAR(50)", "FOREIGN KEY", "Mã khách du lịch sở hữu", "Tham chiếu tới User(ID_User)"],
                ["total_cost", "DECIMAL(12,2)", "NOT NULL", "Tổng chi phí tự tính toán", "Cộng dồn từ chi phí của tất cả các hoạt động lẻ"]
            ]
        },
        {
            "name": "15. UserSchedule (Lịch trình đã lưu của người dùng)",
            "desc": "Bảng liên kết danh sách các lịch trình yêu thích/đã lưu của khách du lịch.",
            "fields": [
                ["ID_User", "VARCHAR(50)", "PRIMARY KEY, FOREIGN KEY", "Mã người dùng lưu", "Tham chiếu tới User(ID_User)"],
                ["ID_Schedule", "VARCHAR(50)", "PRIMARY KEY, FOREIGN KEY", "Mã lịch trình được lưu", "Tham chiếu tới Schedule(ID_Schedule)"]
            ]
        },
        {
            "name": "16. GreenService (Dịch vụ xanh)",
            "desc": "Các dịch vụ thân thiện với môi trường (nhà hàng, khách sạn, xe điện, điểm check-in) do Provider đăng tải.",
            "fields": [
                ["ID_Service", "VARCHAR(50)", "PRIMARY KEY", "Mã định danh dịch vụ", "srv_ + số thứ tự tự tăng hoặc srv_duong_..."],
                ["ID_Vender", "VARCHAR(50)", "FOREIGN KEY", "Nhà cung cấp dịch vụ đại diện", "Tham chiếu tới Vender(ID_Vender)"],
                ["name_service", "VARCHAR(100)", "NOT NULL", "Tên chi tiết dịch vụ xanh", ""],
                ["type", "VARCHAR(20)", "NOT NULL", "Loại dịch vụ xanh", "stay (lưu trú), food (ăn uống), transport (di chuyển), attraction (vui chơi)"],
                ["cost", "DECIMAL(12,2)", "NOT NULL", "Đơn giá của dịch vụ áp dụng", "Đơn vị VNĐ"],
                ["destination", "VARCHAR(100)", "NOT NULL", "Địa điểm cung cấp dịch vụ", "Tỉnh thành hoạt động"],
                ["carbon", "DECIMAL(10,2)", "NOT NULL", "Chỉ số giảm phát thải carbon", "Đơn vị kg CO2, lượng carbon giảm thiểu được"],
                ["image_url", "VARCHAR(255)", "NULL", "Đường dẫn ảnh minh họa", ""],
                ["rating", "DECIMAL(3,2)", "NOT NULL", "Điểm đánh giá trung bình", "Mặc định 5.0"],
                ["bookings_count", "INT", "NOT NULL", "Tổng số lượt đặt thành công", "Mặc định 0"],
                ["current_data", "TEXT/JSONB", "NULL", "Dữ liệu cấu trúc tọa độ bản đồ", "Chứa thông tin kinh độ/vĩ độ (lat,lng) của điểm dịch vụ"]
            ]
        },
        {
            "name": "17. BadgeService (Nhãn dán dịch vụ)",
            "desc": "Bảng liên kết thể hiện các nhãn sinh thái được gán cho một dịch vụ xanh.",
            "fields": [
                ["ID_Badge", "VARCHAR(50)", "PRIMARY KEY, FOREIGN KEY", "Mã nhãn dán liên kết", "Tham chiếu tới Badge(ID_Badge)"],
                ["ID_Service", "VARCHAR(50)", "PRIMARY KEY, FOREIGN KEY", "Mã dịch vụ xanh", "Tham chiếu tới GreenService(ID_Service)"]
            ]
        },
        {
            "name": "18. ServiceBooking (Đặt dịch vụ lẻ)",
            "desc": "Ghi nhận các yêu cầu đặt chỗ dịch vụ xanh từ khách du lịch kèm trạng thái ký quỹ.",
            "fields": [
                ["ID_Booking", "VARCHAR(50)", "PRIMARY KEY", "Mã đặt chỗ dịch vụ lẻ", "BKG_D_ + số thứ tự (ví dụ: BKG_D_4)"],
                ["ID_User", "VARCHAR(50)", "FOREIGN KEY", "Khách hàng đặt dịch vụ", "Tham chiếu tới User(ID_User)"],
                ["ID_Service", "VARCHAR(50)", "FOREIGN KEY", "Dịch vụ xanh được lựa chọn", "Tham chiếu tới GreenService(ID_Service)"],
                ["fullname", "VARCHAR(100)", "NOT NULL", "Họ tên khách đặt dịch vụ", ""],
                ["name_service", "VARCHAR(100)", "NOT NULL", "Tên dịch vụ tại thời điểm đặt", "Lưu vết để tránh biến động giá/tên dịch vụ gốc"],
                ["booking_date", "VARCHAR(20)", "NOT NULL", "Ngày đặt chỗ thực tế", "Định dạng YYYY-MM-DD"],
                ["guests", "INT", "NOT NULL", "Số lượng khách tham gia", "Tối thiểu 1 người"],
                ["value", "DECIMAL(12,2)", "NOT NULL", "Giá trị thanh toán thực tế", "Chốt giá cố định tại thời điểm đặt"],
                ["status", "VARCHAR(20)", "NOT NULL", "Trạng thái đặt đơn lẻ", "pending (đang chờ), deposit (đặt cọc), completed (hoàn thành), rejected (bị từ chối)"],
                ["votes_count", "INT", "NOT NULL", "Số lượt đánh giá đơn đặt chỗ", "Mặc định 0"],
                ["evoucher_code", "VARCHAR(50)", "NULL", "Mã eVoucher phát hành cho khách", "Định dạng EV-vender-..."],
                ["escrow_status", "VARCHAR(20)", "NOT NULL", "Trạng thái ký quỹ tiền ví", "holding (đang giữ), released (đã thanh toán cho đối tác), refunded (đã hoàn tiền cho khách)"]
            ]
        },
        {
            "name": "19. TourBooking (Đặt tour trọn gói)",
            "desc": "Ghi nhận yêu cầu đặt chỗ các tour du lịch trọn gói từ khách du lịch kèm trạng thái ký quỹ.",
            "fields": [
                ["ID_TourBooking", "VARCHAR(50)", "PRIMARY KEY", "Mã đặt tour trọn gói", "BKGT_ + số thứ tự"],
                ["ID_User", "VARCHAR(50)", "FOREIGN KEY", "Khách hàng đặt tour", "Tham chiếu tới User(ID_User)"],
                ["ID_Schedule", "VARCHAR(50)", "FOREIGN KEY", "Tour mẫu được lựa chọn", "Tham chiếu tới Schedule(ID_Schedule)"],
                ["fullname", "VARCHAR(100)", "NOT NULL", "Tên hiển thị khách đặt hàng", ""],
                ["tour_name", "VARCHAR(100)", "NOT NULL", "Tên tour tại thời điểm đặt", "Lưu vết để đối soát độc lập"],
                ["booking_date", "VARCHAR(20)", "NOT NULL", "Ngày đặt tour thực tế", "Định dạng YYYY-MM-DD"],
                ["guests", "INT", "NOT NULL", "Số lượng khách tham gia", "Tối thiểu 1 người"],
                ["value", "DECIMAL(12,2)", "NOT NULL", "Giá trị thanh toán thực tế", "Đơn vị VNĐ"],
                ["status", "VARCHAR(20)", "NOT NULL", "Trạng thái đơn hàng đặt tour", "pending, deposit, completed, rejected"],
                ["votes_count", "INT", "NOT NULL", "Số lượt đánh giá", "Mặc định 0"],
                ["evoucher_code", "VARCHAR(50)", "NULL", "Mã eVoucher phát hành", ""],
                ["escrow_status", "VARCHAR(20)", "NOT NULL", "Trạng thái ký quỹ tiền", "holding (ký quỹ), released (đã đối soát đối tác), refunded (đã hoàn trả khách)"]
            ]
        },
        {
            "name": "20. WithdrawalRequest (Yêu cầu rút tiền đối tác)",
            "desc": "Ghi nhận yêu cầu rút tiền từ số dư ví của nhà cung cấp đối tác về tài khoản ngân hàng.",
            "fields": [
                ["ID_Withdrawal", "VARCHAR(50)", "PRIMARY KEY", "Mã yêu cầu rút tiền đối tác", "WDR_ + số thứ tự"],
                ["ID_User", "VARCHAR(50)", "FOREIGN KEY", "Mã đối tác gửi yêu cầu", "Tham chiếu tới User(ID_User)"],
                ["amount", "DECIMAL(12,2)", "NOT NULL", "Số tiền yêu cầu rút", "Đơn vị VNĐ, không vượt quá số dư ví"],
                ["bank_name", "VARCHAR(50)", "NOT NULL", "Tên ngân hàng nhận tiền", "Ví dụ: Vietcombank, Techcombank"],
                ["account_number", "VARCHAR(30)", "NOT NULL", "Số tài khoản ngân hàng", ""],
                ["account_holder", "VARCHAR(100)", "NOT NULL", "Tên chủ tài khoản nhận", "Họ tên viết hoa không dấu"],
                ["status", "VARCHAR(20)", "NOT NULL", "Trạng thái xử lý yêu cầu", "pending (đang chờ), approved (đã duyệt chuyển khoản), rejected (bị từ chối)"]
            ]
        },
        {
            "name": "21. ScheduleActivity (Hoạt động chi tiết trong lịch trình)",
            "desc": "Các hoạt động cụ thể phân bổ theo từng khung giờ trong các ngày thuộc lịch trình.",
            "fields": [
                ["ID_Activity", "VARCHAR(50)", "PRIMARY KEY", "Mã hoạt động chi tiết", "act_ + số thứ tự tự tăng"],
                ["ID_Service", "VARCHAR(50)", "FOREIGN KEY", "Dịch vụ xanh liên kết", "Tham chiếu tới GreenService(ID_Service), có thể NULL"],
                ["ID_Schedule", "VARCHAR(50)", "FOREIGN KEY", "Mã lịch trình chứa hoạt động", "Tham chiếu tới Schedule(ID_Schedule)"],
                ["day_number", "INT", "NOT NULL", "Số thứ tự ngày của hoạt động", "Giá trị bắt đầu từ 1 đến tổng số ngày của tour"],
                ["time", "VARCHAR(10)", "NOT NULL", "Thời gian diễn ra hoạt động", "Định dạng HH:MM (ví dụ: '14:00')"],
                ["activity_name", "VARCHAR(255)", "NOT NULL", "Tên mô tả hoạt động", ""],
                ["activity_cost", "DECIMAL(12,2)", "NOT NULL", "Chi phí riêng của hoạt động", "Tự động cộng dồn vào tổng tiền lịch trình. Đơn vị VNĐ"],
                ["carbon", "DECIMAL(10,2)", "NOT NULL", "Lượng carbon phát thải hoạt động", "Đơn vị kg CO2"],
                ["icon", "VARCHAR(50)", "NOT NULL", "Biểu tượng đại diện hiển thị", "Tên lớp icon của Bootstrap Icons"],
                ["type", "VARCHAR(30)", "NOT NULL", "Loại hình hoạt động cụ thể", "lodging (nơi ở), food (ăn uống), transport (di chuyển), attraction (điểm đi), experience (trải nghiệm)"],
                ["coordinates", "VARCHAR(100)", "NULL", "Tọa độ bản đồ địa điểm", "Định dạng 'vĩ độ, kinh độ' (latitude, longitude)"]
            ]
        },
        {
            "name": "22. AdCampaign (Chiến dịch quảng cáo)",
            "desc": "Đơn đặt hàng quảng cáo dịch vụ xanh của Provider giúp tăng tiếp cận trên ứng dụng.",
            "fields": [
                ["ID_Campaign", "VARCHAR(50)", "PRIMARY KEY", "Mã định danh chiến dịch", "ADC_ + số thứ tự"],
                ["ID_User", "VARCHAR(50)", "FOREIGN KEY", "Mã nhà quảng cáo thực hiện", "Tham chiếu tới User(ID_User)"],
                ["ID_Service", "VARCHAR(50)", "FOREIGN KEY", "Dịch vụ xanh được quảng bá", "Tham chiếu tới GreenService(ID_Service)"],
                ["campaigns_type", "VARCHAR(50)", "NOT NULL", "Hình thức hiển thị quảng cáo", "Ví dụ: Banner (ảnh lớn), Featured (nổi bật)"],
                ["campaigns_name", "VARCHAR(100)", "NOT NULL", "Tên chương trình chiến dịch", ""],
                ["campaigns_cost", "DECIMAL(12,2)", "NOT NULL", "Chi phí chi trả quảng cáo", "Đơn vị VNĐ"],
                ["duration_days", "INT", "NOT NULL", "Số ngày chạy quảng cáo", ""],
                ["start_date", "VARCHAR(20)", "NOT NULL", "Ngày bắt đầu chạy", "Định dạng YYYY-MM-DD"],
                ["end_date", "VARCHAR(20)", "NOT NULL", "Ngày kết thúc chiến dịch", "Định dạng YYYY-MM-DD"],
                ["status", "VARCHAR(20)", "NOT NULL", "Trạng thái hiển thị chiến dịch", "pending (chờ chạy), active (đang chạy), completed (đã hết hạn)"]
            ]
        },
        {
            "name": "23. CommunityPost (Bài viết cộng đồng)",
            "desc": "Chia sẻ kinh nghiệm du lịch xanh và đánh giá các chuyến đi thực tế của người dùng.",
            "fields": [
                ["ID_Post", "VARCHAR(50)", "PRIMARY KEY", "Mã bài viết cộng đồng", "CUP_ + số thứ tự"],
                ["ID_User", "VARCHAR(50)", "FOREIGN KEY", "Mã người dùng viết bài chia sẻ", "Tham chiếu tới User(ID_User)"],
                ["rating", "DECIMAL(3,2)", "NOT NULL", "Số điểm đánh giá trải nghiệm", "Giá trị số thực từ 1 đến 5 sao"],
                ["text", "TEXT", "NOT NULL", "Nội dung bài viết chia sẻ", "Nội dung văn bản mô tả"],
                ["tour_name", "VARCHAR(100)", "NULL", "Tên hành trình đã đi", ""],
                ["destination", "VARCHAR(100)", "NULL", "Điểm đến của chuyến đi", "Tỉnh thành"],
                ["image_url", "VARCHAR(255)", "NULL", "Đường dẫn ảnh đính kèm bài", ""],
                ["tour_description", "TEXT", "NULL", "Mô tả ngắn gọn về chuyến đi", ""],
                ["days", "INT", "NOT NULL", "Số ngày thực hiện chuyến đi", ""],
                ["likes_count", "INT", "NOT NULL", "Tổng số lượt yêu thích bài đăng", "Mặc định 0"],
                ["comments_count", "INT", "NOT NULL", "Tổng số lượt bình luận bài viết", "Mặc định 0"],
                ["current_data", "TEXT/JSONB", "NULL", "Các thuộc tính bổ sung lưu trữ", "Dùng cho lưu trữ linh hoạt tùy chọn"]
            ]
        },
        {
            "name": "24. CommentPost (Bình luận bài đăng)",
            "desc": "Hệ thống bình luận hỗ trợ lồng ghép phân cấp dưới các bài đăng cộng đồng.",
            "fields": [
                ["ID_Comment", "VARCHAR(50)", "PRIMARY KEY", "Mã bình luận của bài viết", "CEP_ + số thứ tự"],
                ["ID_User", "VARCHAR(50)", "FOREIGN KEY", "Mã người thực hiện bình luận", "Tham chiếu tới User(ID_User)"],
                ["ID_Post", "VARCHAR(50)", "FOREIGN KEY", "Mã bài viết được bình luận", "Tham chiếu tới CommunityPost(ID_Post), có thể NULL"],
                ["ID_Tour", "VARCHAR(50)", "FOREIGN KEY", "Mã tour được bình luận trực tiếp", "Tham chiếu tới Schedule(ID_Schedule), có thể NULL"],
                ["ID_Service", "VARCHAR(50)", "FOREIGN KEY", "Mã dịch vụ xanh được bình luận", "Tham chiếu tới GreenService(ID_Service), có thể NULL"],
                ["parent_comment_id", "VARCHAR(50)", "FOREIGN KEY", "Mã bình luận cha cấp trên", "Tham chiếu tự liên kết tới CommentPost(ID_Comment) dùng hiển thị phân cấp lồng nhau, nullable"],
                ["rating", "DECIMAL(3,2)", "NULL", "Điểm đánh giá sao đính kèm", "Giá trị số thực từ 1 đến 5 sao, nullable"],
                ["text", "TEXT", "NOT NULL", "Nội dung văn bản bình luận", ""],
                ["image_url", "VARCHAR(255)", "NULL", "Đường dẫn ảnh đính kèm bình luận", "nullable"]
            ]
        },
        {
            "name": "25. CPSS (Liên kết bình luận - Tour mẫu)",
            "desc": "Bảng liên kết trung gian kết nối các bình luận đánh giá trực tiếp với các gói Preset Tour.",
            "fields": [
                ["ID_Comment", "VARCHAR(50)", "PRIMARY KEY, FOREIGN KEY", "Mã bình luận đánh giá", "Tham chiếu tới CommentPost(ID_Comment)"],
                ["ID_ScheduleSample", "VARCHAR(50)", "PRIMARY KEY, FOREIGN KEY", "Mã gói tour mẫu được đánh giá", "Tham chiếu tới ScheduleSample(ID_Schedule)"]
            ]
        },
        {
            "name": "26. CPGS (Liên kết bình luận - Dịch vụ xanh)",
            "desc": "Bảng liên kết trung gian kết nối các bình luận đánh giá với các Dịch vụ xanh cụ thể.",
            "fields": [
                ["ID_Comment", "VARCHAR(50)", "PRIMARY KEY, FOREIGN KEY", "Mã bình luận đánh giá", "Tham chiếu tới CommentPost(ID_Comment)"],
                ["ID_Service", "VARCHAR(50)", "PRIMARY KEY, FOREIGN KEY", "Mã dịch vụ xanh được đánh giá", "Tham chiếu tới GreenService(ID_Service)"]
            ]
        }
    ]

    for table_info in tables_data:
        p_tbl_name = doc.add_paragraph()
        p_tbl_name.add_run(table_info["name"]).bold = True
        style_paragraph(p_tbl_name, size_pt=14, bold=True, space_before_pt=12, space_after_pt=4)
        
        p_tbl_desc = doc.add_paragraph()
        p_tbl_desc.add_run(f"Mô tả: {table_info['desc']}")
        style_paragraph(p_tbl_desc, size_pt=12, italic=True, space_after_pt=6)
        
        headers = ["Tên Trường", "Kiểu Dữ Liệu", "Ràng Buộc", "Mô Tả Ý Nghĩa", "Quy tắc dữ liệu (nếu có)"]
        add_styled_table(doc, headers, table_info["fields"])

    doc.add_page_break()
    p_concl_title = doc.add_paragraph()
    p_concl_title.add_run("KẾT LUẬN & ĐỊNH HƯỚNG MÔ HÌNH QUAN HỆ (RDBMS)").bold = True
    style_paragraph(p_concl_title, size_pt=14, bold=True, space_before_pt=12, space_after_pt=6)
    
    p_concl = doc.add_paragraph()
    p_concl.add_run(
        "Các bảng trên được thiết kế để liên kết chặt chẽ với nhau thông qua cơ chế khóa ngoại. "
        "Với mô hình RDBMS chuẩn PostgreSQL sử dụng Sequelize ORM:\n"
        "1. Các trường liên kết được khai báo rõ ràng thành CONSTRAINT FOREIGN KEY ràng buộc cứng.\n"
        "2. Thiết lập cơ chế CASCADE DELETE tại định nghĩa cơ sở dữ liệu để tự động hóa nghiệp vụ dọn dẹp dữ liệu con (ví dụ: tự động xóa Wallet khi xóa User).\n"
        "3. Tích hợp các chỉ mục UNIQUE trực tiếp ở mức schema SQL để đảm bảo an toàn dữ liệu tuyệt đối."
    )
    style_paragraph(p_concl)

    output_path = "DatabaseSpecification.docx"
    doc.save(output_path)
    print("DatabaseSpecification.docx created successfully with 5-column format!")

if __name__ == "__main__":
    main()
